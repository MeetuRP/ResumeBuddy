from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..middleware import get_current_user
from ..models import UserModel
from ..database import get_db
from bson import ObjectId
import tempfile
import io

router = APIRouter()


import hashlib

def _merge_accepted_edits(data: dict, accepted_edits: dict) -> dict:
    """
    Recursively (or flatly) replace text in structured data with accepted AI improvements.
    Uses MD5 hashes of original text as keys, then falls back to normalized matching.
    """
    if not accepted_edits:
        return data

    import json
    import re
    import hashlib

    def normalize(text: str) -> str:
        # Lowercase, remove all non-alphanumeric, collapse whitespace
        text = text.lower()
        text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
        return " ".join(text.split())

    # Map normalized original text to its improvements
    norm_edits = {}
    for key, val in accepted_edits.items():
        if isinstance(val, dict) and "original" in val:
            norm_edits[normalize(val["original"])] = val["improved"]

    # Deep copy
    data_str = json.dumps(data)
    merged = json.loads(data_str)
    stats = {"merged": 0}

    def apply_to_item(item):
        if isinstance(item, str):
            # Check for exact hash first (legacy/precise)
            h = hashlib.md5(item.encode('utf-8')).hexdigest()
            if h in accepted_edits:
                stats["merged"] += 1
                return accepted_edits[h]["improved"]
            
            # Fallback to normalized matching
            item_norm = normalize(item)
            if item_norm in norm_edits:
                stats["merged"] += 1
                return norm_edits[item_norm]
            
            # Check if item is a substring or container of an edit (for bullet points)
            for orig_norm, improved in norm_edits.items():
                if len(orig_norm) > 10 and (orig_norm in item_norm or item_norm in orig_norm):
                    stats["merged"] += 1
                    return improved
        elif isinstance(item, list):
            return [apply_to_item(i) for i in item]
        elif isinstance(item, dict):
            return {k: apply_to_item(v) for k, v in item.items()}
        return item

    result = apply_to_item(merged)
    print(f"[Export] Merged {stats['merged']} edits into document data")
    return result, stats["merged"]


def _build_docx_from_structured(data: dict) -> bytes:
    """
    Build a premium, ATS-compliant DOCX.
    """
    doc = Document()
    
    # Set default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- Header ---
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(data.get("name") or "Resume")
    run.bold = True
    run.font.size = Pt(20)

    contact_parts = []
    if data.get("email"): contact_parts.append(data["email"])
    if data.get("phone"): contact_parts.append(data["phone"])
    links = data.get("links", {})
    if links.get("linkedin"): contact_parts.append("LinkedIn")
    if links.get("github"): contact_parts.append("GitHub")
    
    if contact_parts:
        c_p = doc.add_paragraph(" | ".join(contact_parts))
        c_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_p.paragraph_format.space_after = Pt(12)

    # --- Sections ---
    sections = [
        ("PROFESSIONAL SUMMARY", "summary"),
        ("EXPERIENCE", "experience"),
        ("EDUCATION", "education"),
        ("PROJECTS", "projects"),
        ("SKILLS", "skills"),
        ("CERTIFICATIONS", "certifications")
    ]

    for title, key in sections:
        val = data.get(key)
        if not val: continue
        
        # Section Header
        h = doc.add_heading(title, level=1)
        # Style the heading
        h_run = h.runs[0]
        h_run.font.size = Pt(12)
        h_run.font.color.rgb = RGBColor(0x37, 0x30, 0xa3) # Indigo

        if key == "skills":
            doc.add_paragraph(", ".join(val))
        elif isinstance(val, list):
            for item in val:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(str(item))
        else:
            doc.add_paragraph(str(val))

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_pdf_from_structured(data: dict) -> bytes:
    """
    Build a premium, modern PDF using PyMuPDF (fitz).
    """
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    
    y = 60
    margin = 55
    width = page.rect.width - (2 * margin)
    
    def draw_text(text, font_size=11, is_bold=False, align=0, color=(0,0,0)):
        nonlocal y, page
        font = "helv" if not is_bold else "Helvetica-Bold"
        words = str(text).split()
        if not words: return
        
        lines = []
        curr_line = []
        for w in words:
            test_line = " ".join(curr_line + [w])
            if fitz.get_text_length(test_line, fontname=font, fontsize=font_size) < width:
                curr_line.append(w)
            else:
                lines.append(" ".join(curr_line))
                curr_line = [w]
        lines.append(" ".join(curr_line))
        
        for line in lines:
            if y > page.rect.height - margin:
                page = doc.new_page()
                y = margin
            
            text_w = fitz.get_text_length(line, fontname=font, fontsize=font_size)
            x = margin
            if align == 1: # Center
                x = margin + (width - text_w) / 2
            elif align == 2: # Right
                x = margin + width - text_w
                
            page.insert_text((x, y), line, fontname=font, fontsize=font_size, color=color)
            y += font_size * 1.5
        
    # --- Name ---
    draw_text(data.get("name") or "Resume", font_size=22, is_bold=True, align=1)
    y += 5
    
    # --- Contact ---
    contact = []
    if data.get("email"): contact.append(data["email"])
    if data.get("phone"): contact.append(data["phone"])
    links = data.get("links", {})
    if links.get("linkedin"): contact.append("LinkedIn")
    if links.get("github"): contact.append("GitHub")
    
    if contact:
        draw_text(" | ".join(contact), font_size=10, align=1, color=(0.4, 0.4, 0.4))
    
    y += 25
    
    sections = [
        ("PROFESSIONAL SUMMARY", "summary"),
        ("EXPERIENCE", "experience"),
        ("EDUCATION", "education"),
        ("PROJECTS", "projects"),
        ("SKILLS", "skills"),
        ("CERTIFICATIONS", "certifications")
    ]
    
    for title, key in sections:
        val = data.get(key)
        if not val: continue
        
        # Section Header with Line
        page.draw_line((margin, y - 12), (margin + width, y - 12), color=(0.2, 0.2, 0.6), width=1.5)
        draw_text(title, font_size=12, is_bold=True, color=(0.2, 0.2, 0.6))
        y += 8
        
        if key == "skills":
            draw_text(", ".join(val), font_size=10.5)
        elif isinstance(val, list):
            for item in val:
                # Bullet char (dot)
                page.insert_text((margin + 2, y), "ΓÇó", fontname="helv", fontsize=10)
                # Offset text
                old_margin = margin
                old_width = width
                margin += 15
                width -= 15
                draw_text(item, font_size=10.5)
                margin = old_margin
                width = old_width
                y += 2
        else:
            draw_text(val, font_size=10.5)
        
        y += 18

    return doc.tobytes()


@router.get("/pdf/{resume_id}")
async def export_pdf(resume_id: str, current_user: UserModel = Depends(get_current_user)):
    """
    Generates a clean PDF from the structured data, including all accepted AI changes.
    """
    db = get_db()
    resume = await db.resumes.find_one({
        "_id": ObjectId(resume_id),
        "user_id": str(current_user.id)
    })

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # Merge accepted edits into the base data
    base_data = (
        resume.get("optimized_resume_json")
        or resume.get("structured_resume_json")
        or resume.get("extracted_data", {})
    )
    
    accepted_edits = resume.get("accepted_edits", {})
    merged_data, count = _merge_accepted_edits(base_data, accepted_edits)
    print(f"[Export] Merged {count} edits into PDF data for resume {resume_id}")

    pdf_bytes = _build_pdf_from_structured(merged_data)

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=resume_improved.pdf"}
    )


@router.get("/docx/{resume_id}")
async def export_docx(resume_id: str, current_user: UserModel = Depends(get_current_user)):
    """
    Exports an ATS-compliant DOCX built from the structured resume JSON
    (with any accepted AI optimizations).
    """
    db = get_db()
    resume = await db.resumes.find_one({
        "_id": ObjectId(resume_id),
        "user_id": str(current_user.id)
    })

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    base_data = (
        resume.get("optimized_resume_json")
        or resume.get("structured_resume_json")
        or resume.get("extracted_data", {})
    )
    
    accepted_edits = resume.get("accepted_edits", {})
    merged_data = _merge_accepted_edits(base_data, accepted_edits)

    docx_bytes = _build_docx_from_structured(merged_data)

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=improved_resume.docx"}
    )

